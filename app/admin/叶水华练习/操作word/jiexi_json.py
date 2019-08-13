import json

from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pyecharts.render import make_snapshot
from snapshot_selenium import snapshot
import time
from docx import Document
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




def base_html():
    html = '''<!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Awesome-pyecharts</title>
                <script type="text/javascript" src="echarts.js"></script>
    </head>
    <body>
        <div id="703a9c49431942aebad1569993e9ec62" style="width:900px; height:500px;"></div>
        <script>
            var chart_703a9c49431942aebad1569993e9ec62 = echarts.init(
                document.getElementById('703a9c49431942aebad1569993e9ec62'), 'light', {renderer: 'canvas'});
            var option_703a9c49431942aebad1569993e9ec62 = 配置;
            console.log(option_703a9c49431942aebad1569993e9ec62);
            chart_703a9c49431942aebad1569993e9ec62.setOption(option_703a9c49431942aebad1569993e9ec62);
        </script>
    </body>
    </html>'''

    return html




def jiexi(data):
    with open(data,'r',encoding='utf-8') as file:
        j=json.loads(file.read())
        index=1
        for i in j:
            if i.get('type')=='chart':
                if i.get('option'):
                    option=i.get('option')
                    print(option)

                    ht = str(index) + '.html'
                    with open(ht, 'w') as file:
                        file.write(base_html().replace('配置',json.dumps(option)))
                        file.write('\n')

                    time.sleep(0.2)
                        #filename不允许添加路径，必须在当前目录吓的html文件
                    make_snapshot(snapshot, file_name=ht, output_name="./need/"+str(index)+".png")

                    js_path='./need/'+str(index)+'.js'
                    with open(js_path, 'w', encoding='utf-8') as need:
                        need.write(json.dumps(option, ensure_ascii=False))
                        need.write('\n')

            elif i.get('type')=='text':
                text=i.get('text')
                txt_path='./need/'+str(index)+'.txt'
                with open(txt_path,'w',encoding='utf-8') as need:
                    need.write(text)
                    need.write('\n')

            else:
                return 'error'
            index+=1


def docx(doc_title, doc_title_position):
    docoment = Document()
    section = docoment.sections[0]
    # 文档标题,默认在最左边
    if doc_title:
        # 设置选择标题
        header = section.header
        paragraph = header.paragraphs[0]
        # 如果选择了位置
        if doc_title_position:
            if doc_title_position == 'center':
                paragraph.text = '\t' + doc_title

            if doc_title_position == 'right':
                paragraph.text = '\t\t' + doc_title

            if doc_title_position == 'left':
                paragraph.text = doc_title

        # 默认在最左边
        else:
            paragraph.text = doc_title
    else:
        pass
    docoment.add_heading(doc_title)

    filepath = './need/'
    for i, j, k in os.walk(filepath):
        for name in k:
            houzhui = name.split('.')[-1]
            if houzhui == 'png':
                path=filepath+name
                docoment.add_picture('{}'.format(path),width=Inches(6.5))
            if houzhui == 'txt':
                with open(filepath+name,'r',encoding='utf-8')as f:
                    txt=f.read()
                    docoment.add_paragraph(txt)
            else:
                pass


    #docoment.add_heading('图片的二标题')
    #docoment.add_picture('./need/3.png',width=Inches(6.5))
    #word文档的样式设置

    save_docx_name = './need/first.docx'
    docoment.save(save_docx_name)

def doc_style():
    # 设置默认字体
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    ParagraphStyle1 = document.styles.add_style('ParagraphStyle1', 1)
    # 设置字体尺寸
    ParagraphStyle1.font.size = Pt(30)
    # 设置字体颜色

def ParagrapgStyle1(ti,txt):
    #修改正文的中文字体类型，示例代码：（全局设置）
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    document.add_heading('')
    p=document.add_paragraph('',style='Heading 2')

    #添加了一个段落，不是添加标题，只是在段落上添加文字
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title=p.add_run(ti)
    title.font.name=u'宋体'
    title.bold=True
    title.font.size=Pt(15)
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    p2=document.add_paragraph('')
    #整个段落使用首行缩进0.25厘米
    p2.paragraph_format.first_line_indent = Inches(0.25)
    #txt = '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。'
    text=p2.add_run(txt)
    font=text.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    document.add_picture('1.png',width=Inches(6.3))

    document.add_heading('')

    ti='第二部分'
    p3 = document.add_paragraph('', style='Heading 2')

    # 添加了一个段落，不是添加标题，只是在段落上添加文字
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = p3.add_run(ti)
    title.font.name = u'宋体'
    title.bold = True
    title.font.size = Pt(15)
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    document.save('段落.docx')


def ces(doc_title, doc_title_position, save_docx_name):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    '''
    section = document.sections[0]
    # 文档标题,默认在最左边
    if doc_title:
        # 设置选择标题
        header = section.header
        paragraph = header.paragraphs[0]

        # 如果选择了位置
        if doc_title_position:
            if doc_title_position == 'center':
                paragraph.text = '\t' + doc_title

            if doc_title_position == 'right':
                paragraph.text = '\t\t' + doc_title

            if doc_title_position == 'left':
                paragraph.text = doc_title

        # 默认在最左边
        else:
            paragraph.text = doc_title
    else:
        pass
    '''
    '''
    document.add_heading('')
    p = document.add_paragraph('', style='Heading 2')
    # 添加了一个段落，不是添加标题，只是在段落上添加文字
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = p.add_run(doc_title)
    title.font.name = u'宋体'
    title.bold = True
    title.font.size = Pt(15)
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
'''
    filepath = './need/'
    for i, j, k in os.walk(filepath):

        for name in k:
            houzhui = name.split('.')[-1]
            if houzhui == 'png':

                document.add_heading('')
                p = document.add_paragraph('', style='Heading 2')
                # 添加了一个段落，不是添加标题，只是在段落上添加文字
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title = p.add_run(doc_title)
                title.font.name = u'宋体'
                title.bold = True
                title.font.size = Pt(15)
                title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

                path = filepath + name
                document.add_picture('{}'.format(path), width=Inches(6.5))

            if houzhui == 'txt':
                with open(filepath + name, 'r', encoding='utf-8')as f:
                    txt = f.read()
                    p2 = document.add_paragraph('')
                    # 整个段落使用首行缩进0.25厘米
                    p2.paragraph_format.first_line_indent = Inches(0.25)
                    # txt = '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。'
                    text = p2.add_run(txt)
                    font = text.font
                    font.name = 'Calibri'
                    font.size = Pt(10.5)


            else:
                pass

    document.save(save_docx_name)


if __name__=='__main__':
    #data = 'a.json'
    #jiexi(data)
    #docx('第一张数据报告','left')
    #ti='第一部分'
    #txt = '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。'
    #ParagrapgStyle1(ti,txt)

    ces('第一张数据报告','','theone.docx')
