
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

from docx import Document

# 新建document对象
document=Document()


paragraph=document.add_paragraph("hello world".title())
paragraph.insert_paragraph_before('Python')

document.add_heading('this is a default heading')

document.add_page_break()

table=document.add_table(rows=2,cols=3)
table.add_row()

for row in table.rows:
    for cell in row.cells:
        cell.text='ysh'
cell=table.cell(0,0)

document.save('./need/text.docx')






