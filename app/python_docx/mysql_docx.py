import os

import pymysql
import json
import random
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import pyecharts
from pyecharts.charts import Bar
from snapshot_selenium import snapshot
from pyecharts.render import make_snapshot
import time
import os

from . import  logger
'''
连接mysql数据库，将数据传入到word文档里面
'''


class Mysql_docx:
    def __init__(self):
        if not os.path.exists(self.tmpdir):
            logger.debug('建立输出文件夹:%s',self.tmpdir)
            os.makedirs(self.tmpdir)

    tmpdir = '../report/tmp'
    def base_html(self):
        html = '''<!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Awesome-pyecharts</title>
                    <script type="text/javascript" src="echarts.js"></script>

        </head>
        <body>
            <div id="703a9c49431942aebad1569993e9ec62" style="width:900px; height:500px;"></div>
            <script>
                var chart_703a9c49431942aebad1569993e9ec62 = echarts.init(
                    document.getElementById('703a9c49431942aebad1569993e9ec62'), 'light', {renderer: 'canvas'});
                var option_703a9c49431942aebad1569993e9ec62 = 配置;
                console.log(option_703a9c49431942aebad1569993e9ec62);
                chart_703a9c49431942aebad1569993e9ec62.setOption(option_703a9c49431942aebad1569993e9ec62);
            </script>
        </body>
        </html>'''

        return html

    def jiexi(self,data,doc_title,filename):
        #with open(data, 'r', encoding='utf-8') as file:
        j = json.loads(data)
        index = 1
        for i in j:
            if i.get('type') == 'chart':
                if i.get('option'):
                    option = i.get('option')
                    print(option)

                    ht =  os.path.join(self.tmpdir, str(index) + '.html')
                    logger.debug('生成html:%s,路径:%s',ht,os.path.abspath(ht))

                    with open(ht, 'w') as file:
                        file.write(self.base_html().replace('配置', json.dumps(option)))
                        file.write('\n')


                    # time.sleep(0.2)
                    # filename不允许添加路径，必须在当前目录吓的html文件
                    output_name =  os.path.join(self.tmpdir,  str(index) + ".png")
                    self.html_echarts(ht,output_name)

                    js_path =  os.path.join(self.tmpdir, str(index) + '.js')
                    self.save_js(option,js_path)

            elif i.get('type') == 'text':
                text = i.get('text')
                txt_path =  os.path.join(self.tmpdir, str(index) + '.txt')
                self.save_txt(text,txt_path)

            else:
                continue
            index += 1


        #word文档

        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        document.add_heading('')
        p = document.add_paragraph('', style='Heading 2')
        # 添加了一个段落，不是添加标题，只是在段落上添加文字
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title = p.add_run(doc_title)
        title.font.name = u'宋体'
        title.bold = True
        title.font.size = Pt(15)
        title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        filepath = self.tmpdir
        for i, j, k in os.walk(filepath):
            logger.debug('docx文件列表:%s',k)
            k.sort()
            for name in  k:

                houzhui = name.split('.')[-1]
                if houzhui == 'png':
                    path = os.path.join( filepath , name)
                    document.add_picture('{}'.format(path), width=Inches(6.5))
                if houzhui == 'txt':
                    with open(os.path.join( filepath , name), 'r', encoding='utf-8')as f:
                        txt = f.read()
                        p2 = document.add_paragraph('')
                        # 整个段落使用首行缩进0.25厘米
                        p2.paragraph_format.first_line_indent = Inches(0.25)
                        # txt = '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。'
                        text = p2.add_run(txt)
                        font = text.font
                        font.name = 'Calibri'
                        font.size = Pt(10.5)

                else:
                    pass
        logger.debug('生成路径:%s,标题:%s',filename,title)
        # save_docx_name= os.path.join(self.outdir,filename)
        document.save(filename)
        return  filename


    # 传入HTML，将其转为图片
    def html_echarts(self, html, output_name):

        # 选择要保存的图片的名字,需要添加后缀

        # filename不允许添加路径，必须在当前目录吓的html文件

        make_snapshot(snapshot, html, output_name)


    def save_txt(self,text,txt_path):
        with open(txt_path, 'w', encoding='utf-8') as need:
            need.write(text)
            need.write('\n')
        need.close()

    #保存js
    def save_js(self, option, js_path):
        with open(js_path, 'w', encoding='utf-8') as need:
            need.write(json.dumps(option, ensure_ascii=False))
            need.write('\n')
        need.close()

    # 传入text和imgs上的图片，将其作用在docx上，生成word文档
    def docx(self, doc_title, doc_title_position, txt,picture_path,save_docx_name):
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        section = document.sections[0]
        # 文档标题,默认在最左边
        if doc_title:
            # 设置选择标题
            header = section.header
            paragraph = header.paragraphs[0]

            # 如果选择了位置
            if doc_title_position:
                if doc_title_position == 'center':
                    paragraph.text = '\t' + doc_title

                if doc_title_position == 'right':
                    paragraph.text = '\t\t' + doc_title

                if doc_title_position == 'left':
                    paragraph.text = doc_title

            # 默认在最左边
            else:
                paragraph.text = doc_title
        else:
            pass
        document.add_heading('')
        p = document.add_paragraph('', style='Heading 2')
        # 添加了一个段落，不是添加标题，只是在段落上添加文字
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title=p.add_run(doc_title)
        title.font.name = u'宋体'
        title.bold = True
        title.font.size = Pt(15)
        title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        filepath = 'need/'
        for i, j, k in os.walk(filepath):
            for name in k:
                houzhui = name.split('.')[-1]
                if houzhui == 'png':
                    path = os.path.join(filepath  ,name)
                    document.add_picture('{}'.format(path), width=Inches(6.5))
                if houzhui == 'txt':
                    with open(os.path.join(filepath  ,name), 'r', encoding='utf-8')as f:
                        txt = f.read()
                        p2 = document.add_paragraph('')
                        # 整个段落使用首行缩进0.25厘米
                        p2.paragraph_format.first_line_indent = Inches(0.25)
                        # txt = '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。'
                        text = p2.add_run(txt)
                        font = text.font
                        font.name = 'Calibri'
                        font.size = Pt(10.5)

                else:
                    pass

        document.save(save_docx_name)




if __name__ == '__main__':
    MD = Mysql_docx()
    html = MD.jiexi('a.json')
    MD.docx('第一部分',)
